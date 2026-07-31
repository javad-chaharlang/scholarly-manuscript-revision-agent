'''Evaluate deterministic Phase 7 readiness without bypassing human approval.'''

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook

from scholarly_revision.models.release import (
    ConsistencyFinding, FinalReleaseCheck, FinalReleaseChecklist,
    FinalReleaseReport,
)
from scholarly_revision.models.response_package import ResponsePackage, ResponseStatus
from scholarly_revision.models.scientific_audit import (
    AuditIssueStatus, AuditSeverity, FinalReleaseReadiness,
)
from scholarly_revision.services.config_loader import load_project_manifest
from scholarly_revision.services.gap_analysis_service import read_json, write_json
from scholarly_revision.services.manual_visual_qa_service import (
    evaluate_manual_visual_qa,
)
from scholarly_revision.services.response_letter_service import load_response_sources
from scholarly_revision.tools.docx_clean_copy import validate_text_equivalence
from scholarly_revision.tools.highlight_auditor import audit_highlights


RELEASE_CHECK_CATEGORIES = (
    'project metadata complete',
    'every reviewer comment accounted for',
    'every response verified',
    'all required author approvals recorded',
    'manuscript changes applied',
    'clean/highlighted manuscript equivalence verified',
    'highlight policy verified',
    'references structurally verified',
    'empirical results verified',
    'unresolved blockers reviewed',
    'scientific QA readiness acceptable',
    'front matter reviewed',
    'figures and tables checked',
    'equations and symbols checked',
    'response letter verified',
    'final files open successfully',
    'confidential source files excluded from Git',
    'release manifest complete',
    'rendered Word documents visually inspected',
    'final human release approval recorded',
)


def _approval(root: Path, explicit: dict[str, Any] | None) -> dict[str, Any]:
    if explicit is not None:
        return explicit
    path = root / 'audit' / 'final_release_approval.json'
    return read_json(path) if path.is_file() else {}


def _check(category: str, passed: bool, *evidence: str, notes: str | None = None) -> FinalReleaseCheck:
    return FinalReleaseCheck(
        category=category, passed=passed, evidence=list(evidence), notes=notes
    )


def _qa_pass(qa: dict[str, Any], categories: set[str]) -> bool:
    issues = qa.get('issues', []) if isinstance(qa, dict) else []
    return bool(qa) and not any(
        str(item.get('category')) in categories
        and str(item.get('status', 'OPEN')) in {'OPEN', 'ACKNOWLEDGED'}
        and str(item.get('severity')) in {'BLOCKER', 'CRITICAL', 'MAJOR'}
        for item in issues
    )


def evaluate_final_release(
    project_root: str | Path,
    findings: list[ConsistencyFinding | dict] | None = None,
    *,
    final_approval: dict[str, Any] | None = None,
) -> FinalReleaseReport:
    root = Path(project_root).expanduser().resolve()
    sources = load_response_sources(root)
    manifest = load_project_manifest(root / 'config' / 'project_manifest.yaml')
    package = ResponsePackage.model_validate(
        read_json(root / 'working' / 'response_package.json')
    )
    if findings is None:
        report_path = root / 'audit' / 'final_consistency_report.json'
        raw = read_json(report_path) if report_path.is_file() else {'findings': []}
        findings = raw.get('findings', [])
    consistency = [ConsistencyFinding.model_validate(item) for item in findings]
    qa = sources['qa']
    approval = _approval(root, final_approval)
    approved = approval.get('approved') is True or approval.get('decision') == 'APPROVE'
    approver = approval.get('decision_maker') or approval.get('approved_by')
    approved_at = approval.get('decision_timestamp') or approval.get('approved_at')
    highlighted = root / 'outputs' / 'Revised_Manuscript_Highlighted.docx'
    clean = root / 'outputs' / 'Revised_Manuscript_Clean.docx'
    response_docx = root / 'outputs' / 'Response_to_Reviewers.docx'
    workbook = root / 'outputs' / 'Revision_Master.xlsx'
    qa_workbook = root / 'outputs' / 'Scientific_QA_Report.xlsx'
    comments = sources['comments']
    entry_ids = {item.comment_id for item in package.entries}
    comment_ids = {item.comment_id for item in comments}
    blocking = [
        item for item in consistency if item.status in {
            AuditIssueStatus.OPEN, AuditIssueStatus.ACKNOWLEDGED,
        } and item.severity in {AuditSeverity.BLOCKER, AuditSeverity.CRITICAL}
    ]
    qa_blockers = [
        item for item in qa.get('issues', []) if
        str(item.get('status', 'OPEN')) in {'OPEN', 'ACKNOWLEDGED'}
        and str(item.get('severity')) in {'BLOCKER', 'CRITICAL'}
    ] if isinstance(qa, dict) else []
    change_actions = {item.action_id for item in sources['changes']}
    approved_actions = {
        item.action_id for item in sources['actions']
        if item.approval_state.value == 'APPROVED'
    }
    try:
        equivalent = (
            highlighted.is_file() and clean.is_file()
            and validate_text_equivalence(highlighted, clean)
        )
    except Exception:
        equivalent = False
    highlight_ok = False
    if highlighted.is_file() and clean.is_file():
        try:
            highlight_result = audit_highlights(
                highlighted, clean,
                change_log=root / 'audit' / 'change_log.json'
                if (root / 'audit' / 'change_log.json').is_file() else None,
                reference_registry=sources['references'],
            )
            highlight_ok = not any(
                item.severity in {AuditSeverity.BLOCKER, AuditSeverity.CRITICAL}
                and item.status in {AuditIssueStatus.OPEN, AuditIssueStatus.ACKNOWLEDGED}
                for item in highlight_result.issues
            )
        except Exception:
            highlight_ok = False
    results = sources.get('results', [])
    empirical_ok = all(
        str(item.get('result_status')) not in {'FINAL', 'VERIFIED'}
        or str(item.get('evidence_status')) == 'VERIFIED'
        for item in results
    )
    response_report_path = root / 'audit' / 'response_verification_report.json'
    response_report = read_json(response_report_path) if response_report_path.is_file() else {}
    visual_path = root / 'audit' / 'visual_inspection.json'
    visual_evaluation = evaluate_manual_visual_qa(root)
    visual_complete = visual_evaluation.passed
    visual = {
        'schema_version': 1,
        'status': (
            'APPROVED' if visual_complete else 'MANUAL_VISUAL_QA_REQUIRED'
        ),
        'passed': visual_complete,
        'scope': [
            'Response_to_Reviewers.docx',
            'Revised_Manuscript_Highlighted.docx',
            'Revised_Manuscript_Clean.docx',
            'Revision_Master.xlsx',
            'Scientific_QA_Report.xlsx',
        ],
        'decision_record': str(visual_evaluation.record_path),
        'notes': visual_evaluation.reason,
    }
    write_json(visual_path, visual)
    files_open = True
    try:
        Document(highlighted)
        Document(clean)
        Document(response_docx)
        load_workbook(workbook, read_only=True).close()
        load_workbook(qa_workbook, read_only=True).close()
    except Exception:
        files_open = False
    repository_root = Path(__file__).resolve().parents[3]
    outside_git = repository_root not in root.parents and root != repository_root
    metadata_ok = all([
        manifest.manuscript_title != 'UNSPECIFIED',
        manifest.manuscript_id != 'UNSPECIFIED',
        manifest.journal != 'UNSPECIFIED',
        manifest.revision_round >= 1,
    ])
    checks = [
        _check(RELEASE_CHECK_CATEGORIES[0], metadata_ok),
        _check(RELEASE_CHECK_CATEGORIES[1], entry_ids == comment_ids),
        _check(RELEASE_CHECK_CATEGORIES[2], bool(package.entries) and all(
            item.response_status is ResponseStatus.VERIFIED for item in package.entries
        )),
        _check(RELEASE_CHECK_CATEGORIES[3], all(
            item.author_approved for item in package.entries
        )),
        _check(RELEASE_CHECK_CATEGORIES[4], approved_actions.issubset(change_actions)),
        _check(RELEASE_CHECK_CATEGORIES[5], equivalent),
        _check(RELEASE_CHECK_CATEGORIES[6], highlight_ok),
        _check(RELEASE_CHECK_CATEGORIES[7], _qa_pass(qa, {'REFERENCE', 'CITATION'})),
        _check(RELEASE_CHECK_CATEGORIES[8], empirical_ok and _qa_pass(
            qa, {'RESULT_INTEGRITY', 'NUMERICAL_CONSISTENCY'}
        )),
        _check(RELEASE_CHECK_CATEGORIES[9], not blocking and not qa_blockers),
        _check(RELEASE_CHECK_CATEGORIES[10], bool(qa) and str(
            qa.get('final_release_readiness')
        ) in {'READY', 'READY_WITH_WARNINGS'}),
        _check(RELEASE_CHECK_CATEGORIES[11], _qa_pass(qa, {'FRONT_MATTER'})),
        _check(RELEASE_CHECK_CATEGORIES[12], _qa_pass(qa, {'FIGURE_TABLE'})),
        _check(RELEASE_CHECK_CATEGORIES[13], _qa_pass(qa, {'EQUATION_SYMBOL'})),
        _check(RELEASE_CHECK_CATEGORIES[14], response_report.get('passed') is True),
        _check(RELEASE_CHECK_CATEGORIES[15], files_open),
        _check(RELEASE_CHECK_CATEGORIES[16], outside_git),
        _check(RELEASE_CHECK_CATEGORIES[17], all(
            path.is_file() for path in (
                highlighted, clean, response_docx, workbook, qa_workbook,
            )
        )),
        _check(
            RELEASE_CHECK_CATEGORIES[18], visual_complete,
            str(visual_evaluation.record_path),
            notes=None if visual_complete else 'MANUAL_VISUAL_QA_REQUIRED',
        ),
        _check(RELEASE_CHECK_CATEGORIES[19], approved),
    ]
    prohibited_failures = {
        RELEASE_CHECK_CATEGORIES[index] for index in (1, 2, 4, 5, 8, 9)
    }
    failed_required = [item for item in checks if item.required and not item.passed]
    if blocking or qa_blockers or any(
        item.category in prohibited_failures for item in failed_required
    ):
        readiness = FinalReleaseReadiness.BLOCKED
    elif failed_required:
        readiness = FinalReleaseReadiness.NOT_READY
    else:
        open_findings = [
            item for item in consistency
            if item.status in {AuditIssueStatus.OPEN, AuditIssueStatus.ACKNOWLEDGED}
        ]
        if any(item.severity is AuditSeverity.MAJOR for item in open_findings):
            readiness = FinalReleaseReadiness.NOT_READY
        elif open_findings or str(qa.get('final_release_readiness')) == 'READY_WITH_WARNINGS':
            readiness = FinalReleaseReadiness.READY_WITH_WARNINGS
        else:
            readiness = FinalReleaseReadiness.READY
    checklist = FinalReleaseChecklist(
        generated_at=datetime.now(UTC), checks=checks, readiness=readiness
    )
    blocker_reasons = [
        item.description for item in consistency if item.blocking
    ] + [str(item.get('description')) for item in qa_blockers]
    blocker_reasons.extend(
        item.category for item in failed_required
        if item.category in prohibited_failures
    )
    warning_reasons = [
        item.description for item in consistency
        if item.status in {AuditIssueStatus.OPEN, AuditIssueStatus.ACKNOWLEDGED}
        and item.severity in {AuditSeverity.MINOR, AuditSeverity.INFORMATIONAL}
    ]
    if not visual_complete:
        warning_reasons.append('MANUAL_VISUAL_QA_REQUIRED')
    report = FinalReleaseReport(
        generated_at=datetime.now(UTC),
        readiness=readiness,
        checklist=checklist,
        consistency_findings=consistency,
        blocker_reasons=list(dict.fromkeys(blocker_reasons)),
        warning_reasons=list(dict.fromkeys(warning_reasons)),
        final_author_approved=approved,
        final_approval_by=str(approver) if approved and approver else None,
        final_approval_at=approved_at if approved else None,
        release_permitted=(
            readiness is FinalReleaseReadiness.READY
            or readiness is FinalReleaseReadiness.READY_WITH_WARNINGS and approved
        ),
    )
    write_json(
        root / 'audit' / 'final_release_checklist.json',
        checklist.model_dump(mode='json'),
    )
    write_json(
        root / 'audit' / 'final_release_report.json',
        report.model_dump(mode='json'),
    )
    write_json(
        root / 'outputs' / 'Final_Release_Report.json',
        report.model_dump(mode='json'),
    )
    return report


class FinalReleaseService:
    def evaluate(self, project_root: str | Path, **kwargs: Any) -> FinalReleaseReport:
        return evaluate_final_release(project_root, **kwargs)
