'''Build and inspect editable response-to-reviewers DOCX files.'''

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from scholarly_revision.models.enums import HighlightColor
from scholarly_revision.models.response_package import ResponseEntry, ResponsePackage


_COLORS = {
    HighlightColor.YELLOW: WD_COLOR_INDEX.YELLOW,
    HighlightColor.BRIGHT_GREEN: WD_COLOR_INDEX.BRIGHT_GREEN,
    HighlightColor.VIOLET: WD_COLOR_INDEX.VIOLET,
}
_SHADING = {
    HighlightColor.YELLOW: 'FFF2CC',
    HighlightColor.BRIGHT_GREEN: 'E2F0D9',
    HighlightColor.VIOLET: 'E4DFEC',
}
_FIELD_LABELS = (
    'Comment:',
    "Author's response:",
    'Changes made in the manuscript:',
    'Location:',
    'Highlight:',
)


@dataclass(frozen=True, slots=True)
class ResponseDocxEntryRecord:
    '''One response block extracted deterministically from the generated DOCX.'''

    heading: str
    comment: str
    author_response: str
    changes_made: str
    location: str
    highlight: str
    heading_highlight: HighlightColor | None


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        properties.append(shading)
    shading.set(qn('w:fill'), color)


def _set_cell_border(cell, color: str = 'B7B7B7') -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        properties.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '4')
        node.set(qn('w:color'), color)
        borders.append(node)


def _black(document: Document) -> None:
    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for name in ('Title', 'Heading 1', 'Heading 2', 'Heading 3'):
        styles[name].font.name = 'Arial'
        styles[name].font.color.rgb = RGBColor(0, 0, 0)
    if 'Response Label' not in styles:
        label = styles.add_style('Response Label', WD_STYLE_TYPE.PARAGRAPH)
        label.font.name = 'Arial'
        label.font.size = Pt(10)
        label.font.bold = True


def _field(cell, label: str, value: str) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(label + '\n')
    label_run.bold = True
    paragraph.add_run(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _entry_heading(entry: ResponseEntry) -> str:
    if entry.reviewer_source.value == 'REVIEWER':
        return f'Reviewer {entry.reviewer_number}, Comment {entry.sequence_number}'
    if entry.reviewer_source.value == 'EDITOR':
        return f'Editor, Comment {entry.sequence_number}'
    return f'General Comment {entry.sequence_number}'


def _add_entry(document: Document, entry: ResponseEntry) -> None:
    heading = document.add_heading(_entry_heading(entry), level=2)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.highlight_color = _COLORS[entry.highlight]
        run.font.color.rgb = RGBColor(0, 0, 0)
    table = document.add_table(rows=5, cols=1)
    values = (
        ('Comment:', entry.exact_comment, ''),
        ('Author\'s response:', entry.author_response, 'Response pending.'),
        ('Changes made in the manuscript:', entry.changes_made, 'No manuscript change reported.'),
        ('Location:', '; '.join(entry.verified_locations), 'Not required.'),
        ('Highlight:', entry.highlight.value.replace('_', ' ').title(), ''),
    )
    for row, (label, value, fallback) in zip(table.rows, values):
        _field(row.cells[0], label, value or fallback)
        _set_cell_border(row.cells[0])
    _shade(table.rows[0].cells[0], _SHADING[entry.highlight])
    document.add_paragraph()


def build_response_docx(
    package: ResponsePackage | dict,
    output_path: str | Path,
) -> Path:
    '''Build a professional editable DOCX from a strictly validated package.'''

    response = ResponsePackage.model_validate(package)
    destination = Path(output_path).expanduser().resolve()
    if destination.suffix.lower() != '.docx':
        raise ValueError('response letter output must be a DOCX file')
    if destination.exists():
        raise FileExistsError(f'response letter already exists: {destination}')
    document = Document()
    _black(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    title = document.add_heading('Response to the Editor and Reviewers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = document.add_table(rows=4, cols=2)
    metadata = (
        ('Manuscript title', response.manuscript_title),
        ('Manuscript ID', response.manuscript_id),
        ('Journal', response.journal),
        ('Revision round', str(response.revision_round)),
    )
    for row, values in zip(info.rows, metadata):
        row.cells[0].text, row.cells[1].text = values
        row.cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_border(row.cells[0])
        _set_cell_border(row.cells[1])
    document.add_heading('Cover Letter to the Editor', level=1)
    document.add_paragraph(response.cover_letter.salutation)
    for paragraph in response.cover_letter.body_paragraphs:
        document.add_paragraph(paragraph)
    document.add_paragraph(response.cover_letter.closing)
    document.add_heading('Summary of Major Revisions', level=1)
    if response.summary_of_major_revisions:
        for item in response.summary_of_major_revisions:
            document.add_paragraph(item, style='List Bullet')
    else:
        document.add_paragraph('No separate major-revision summary was supplied.')
    for response_section in response.sections:
        document.add_heading(response_section.title, level=1)
        for entry in response_section.entries:
            _add_entry(document, entry)
    document.add_heading('General Revisions', level=1)
    if response.general_revisions:
        for item in response.general_revisions:
            document.add_paragraph(item, style='List Bullet')
    else:
        document.add_paragraph('No additional general revisions are reported.')
    document.add_heading('Closing Statement', level=1)
    document.add_paragraph(response.closing_statement)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    Document(destination)
    return destination


def response_docx_text(path: str | Path) -> str:
    document = Document(Path(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return '\n'.join(parts)


def response_docx_entry_records(
    path: str | Path,
) -> tuple[ResponseDocxEntryRecord, ...]:
    '''Extract ordered response blocks for field-by-field consistency checks.'''

    document = Document(Path(path))
    headings = [
        paragraph for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name == 'Heading 2'
    ]
    blocks: list[dict[str, str]] = []
    for table in document.tables:
        values: dict[str, str] = {}
        for row in table.rows:
            if len(row.cells) != 1:
                continue
            text = row.cells[0].text
            for label in _FIELD_LABELS:
                if text == label or text.startswith(label + '\n'):
                    values[label] = text[len(label):].lstrip('\n')
                    break
        if values:
            if set(values) != set(_FIELD_LABELS):
                raise ValueError('response DOCX contains an incomplete response block')
            blocks.append(values)
    if len(headings) != len(blocks):
        raise ValueError('response DOCX heading and response-block counts differ')
    reverse_colors = {value: key for key, value in _COLORS.items()}
    records = []
    for heading, values in zip(headings, blocks):
        visible = {
            run.font.highlight_color for run in heading.runs
            if run.text and run.font.highlight_color is not None
        }
        visible_color = next(iter(visible)) if len(visible) == 1 else None
        records.append(ResponseDocxEntryRecord(
            heading=heading.text,
            comment=values['Comment:'],
            author_response=values["Author's response:"],
            changes_made=values['Changes made in the manuscript:'],
            location=values['Location:'],
            highlight=values['Highlight:'],
            heading_highlight=reverse_colors.get(visible_color),
        ))
    return tuple(records)
