'''Apply and audit only repository-authorized revision highlights.'''

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from scholarly_revision.models.enums import HighlightColor


SYSTEM_STYLE_PREFIX = 'Scholarly Revision Highlight / '
HIGHLIGHT_INDEX = {
    HighlightColor.YELLOW: WD_COLOR_INDEX.YELLOW,
    HighlightColor.BRIGHT_GREEN: WD_COLOR_INDEX.BRIGHT_GREEN,
    HighlightColor.LIGHT_BLUE: WD_COLOR_INDEX.TURQUOISE,
    HighlightColor.PINK: WD_COLOR_INDEX.PINK,
    HighlightColor.TEAL: WD_COLOR_INDEX.TEAL,
    HighlightColor.DARK_YELLOW: WD_COLOR_INDEX.DARK_YELLOW,
    HighlightColor.GRAY_25: WD_COLOR_INDEX.GRAY_25,
    HighlightColor.DARK_BLUE: WD_COLOR_INDEX.DARK_BLUE,
    HighlightColor.RED: WD_COLOR_INDEX.RED,
    HighlightColor.VIOLET: WD_COLOR_INDEX.VIOLET,
}


def apply_revision_highlight(
    run: Run,
    highlight: HighlightColor | str,
    *,
    change_id: str,
) -> None:
    color = HighlightColor(highlight)
    prior_style = run.style
    prior_style_id = prior_style.style_id if prior_style is not None else 'NONE'
    safe_change_id = change_id.replace('/', '_')
    marker_name = (
        f'{SYSTEM_STYLE_PREFIX}{color.value} / '
        f'{safe_change_id} / {prior_style_id}'
    )
    styles = run.part.document.styles
    try:
        marker_style = styles[marker_name]
    except KeyError:
        marker_style = styles.add_style(marker_name, WD_STYLE_TYPE.CHARACTER)
        marker_style.hidden = True
        if prior_style is not None and prior_style.type == WD_STYLE_TYPE.CHARACTER:
            marker_style.base_style = prior_style
    run.style = marker_style
    run.font.highlight_color = HIGHLIGHT_INDEX[color]


def _system_marker(run: Run) -> tuple[HighlightColor, str] | None:
    style = run.style
    if style is None or not style.name.startswith(SYSTEM_STYLE_PREFIX):
        return None
    parts = style.name[len(SYSTEM_STYLE_PREFIX):].split(' / ', 2)
    if len(parts) != 3:
        return None
    try:
        color = HighlightColor(parts[0])
    except ValueError:
        return None
    return color, parts[1]


def is_system_highlight(run: Run) -> bool:
    return _system_marker(run) is not None


def iter_document_paragraphs(document) -> Iterable[Paragraph]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def remove_revision_highlights(document) -> int:
    removed = 0
    for paragraph in iter_document_paragraphs(document):
        for run in paragraph.runs:
            if not is_system_highlight(run):
                continue
            run.font.highlight_color = None
            marker_style = run.style
            prior_style = marker_style.base_style if marker_style is not None else None
            if prior_style is not None and prior_style.style_id != 'DefaultParagraphFont':
                run.style = prior_style
            else:
                run.style = None
            removed += 1
    return removed


def audit_revision_highlights(path: str | Path) -> dict[str, object]:
    document = Document(Path(path))
    records: list[dict[str, str | None]] = []
    violations: list[str] = []
    for paragraph in iter_document_paragraphs(document):
        for run in paragraph.runs:
            marker = _system_marker(run)
            if marker is None:
                continue
            color, change_id = marker
            actual = run.font.highlight_color
            expected = HIGHLIGHT_INDEX[color]
            if actual != expected:
                violations.append('A system-marked run does not match its policy highlight.')
            records.append({
                'change_id': change_id,
                'highlight': color.value,
                'text_hash_basis': 'run text intentionally omitted from audit output',
            })
    return {
        'system_highlight_count': len(records),
        'records': records,
        'violations': violations,
        'passed': not violations,
    }
