'''Exact-target DOCX mutation for approved plain-content Phase 5 drafts.'''

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from scholarly_revision.models.enums import RevisionOperation
from scholarly_revision.models.revision_draft import RevisionDraft
from scholarly_revision.services.project_workspace import sha256_file
from scholarly_revision.services.revision_drafting_service import sha256_text
from scholarly_revision.tools.docx_highlight_manager import apply_revision_highlight
from scholarly_revision.tools.manuscript_structure_reader import read_manuscript_structure


class RevisionApplicationError(ValueError):
    '''Raised before saving when an exact or safe application condition fails.'''


@dataclass(frozen=True, slots=True)
class AppliedMutation:
    draft_id: str
    action_id: str
    comment_ids: tuple[str, ...]
    operation: str
    target_section: str
    target_element_id: str
    old_text: str
    new_text: str
    highlight: str
    warnings: tuple[str, ...] = ()


_BLOCKED_XML = {
    'oMath': 'equation XML',
    'oMathPara': 'equation XML',
    'fldChar': 'field XML',
    'instrText': 'field or citation XML',
    'fldSimple': 'field XML',
    'drawing': 'drawing or image',
    'pict': 'legacy image',
    'object': 'embedded object',
    'oleObject': 'embedded object',
    'hyperlink': 'hyperlink',
    'bookmarkStart': 'bookmark',
    'bookmarkEnd': 'bookmark',
    'commentRangeStart': 'comment anchor',
    'commentRangeEnd': 'comment anchor',
    'footnoteReference': 'footnote reference',
    'endnoteReference': 'endnote reference',
    'ins': 'tracked-change XML',
    'del': 'tracked-change XML',
    'moveFrom': 'tracked-change XML',
    'moveTo': 'tracked-change XML',
    'sdt': 'content control',
}


def complex_content_reasons(paragraph: Paragraph) -> list[str]:
    reasons: list[str] = []
    for local_name, label in _BLOCKED_XML.items():
        if paragraph._p.xpath(f'.//*[local-name()="{local_name}"]'):
            reasons.append(label)
    return list(dict.fromkeys(reasons))


def _paragraph_map(document) -> dict[str, Paragraph]:
    result: dict[str, Paragraph] = {}
    counter = 0

    def add(paragraph: Paragraph) -> None:
        nonlocal counter
        counter += 1
        result[f'PAR-{counter:04d}'] = paragraph

    for child in document.element.body.iterchildren():
        if child.tag == qn('w:p'):
            add(Paragraph(child, document))
        elif child.tag == qn('w:tbl'):
            table = Table(child, document)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        add(paragraph)
    return result


def _structure_aliases(structure) -> tuple[dict[str, Any], dict[str, int]]:
    aliases: dict[str, Any] = {}
    order: dict[str, int] = {}
    for element in structure.elements:
        for identifier in (element.element_id, element.paragraph_id):
            if identifier:
                aliases[identifier] = element
                order[identifier] = element.order_index
    return aliases, order


def _approved(draft: RevisionDraft) -> bool:
    return (
        draft.approval_state.value == 'APPROVED'
        and draft.author_decision is not None
        and draft.author_decision.value in {
            'APPROVE_TEXT', 'APPROVE_TEXT_WITH_MODIFICATION'
        }
        and draft.draft_status.value == 'APPROVED'
    )


def _validate_conflicts(drafts: list[RevisionDraft]) -> None:
    owners: dict[str, RevisionDraft] = {}
    for draft in drafts:
        for target in draft.target_element_ids:
            previous = owners.get(target)
            if previous is None:
                owners[target] = draft
                continue
            compatible = (
                previous.draft_id in draft.compatible_with_draft_ids
                and draft.draft_id in previous.compatible_with_draft_ids
                and previous.operation in {
                    RevisionOperation.INSERT_BEFORE, RevisionOperation.INSERT_AFTER
                }
                and draft.operation in {
                    RevisionOperation.INSERT_BEFORE, RevisionOperation.INSERT_AFTER
                }
            )
            if not compatible:
                raise RevisionApplicationError(
                    f'overlapping or duplicate target operations: '
                    f'{previous.draft_id}, {draft.draft_id} -> {target}'
                )


def _replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    draft: RevisionDraft,
    change_id: str,
) -> None:
    runs = list(paragraph.runs)
    if runs:
        primary = runs[0]
        primary.text = text
        for run in runs[1:]:
            run._r.getparent().remove(run._r)
    else:
        primary = paragraph.add_run(text)
    apply_revision_highlight(primary, draft.highlight, change_id=change_id)


def _new_sibling(
    paragraph: Paragraph,
    text: str,
    *,
    before: bool,
    draft: RevisionDraft,
    change_id: str,
) -> Paragraph:
    new_p = OxmlElement('w:p')
    if paragraph._p.pPr is not None:
        new_p.append(deepcopy(paragraph._p.pPr))
    if before:
        paragraph._p.addprevious(new_p)
    else:
        paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    run = created.add_run(text)
    apply_revision_highlight(run, draft.highlight, change_id=change_id)
    return created


def _delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise RevisionApplicationError('target paragraph has no document parent')
    parent.remove(paragraph._p)


def apply_docx_revisions(
    source_path: str | Path,
    output_path: str | Path,
    drafts: Iterable[RevisionDraft],
    *,
    expected_source_hash: str,
) -> tuple[AppliedMutation, ...]:
    source = Path(source_path).expanduser().resolve()
    destination = Path(output_path).expanduser().resolve()
    if source == destination:
        raise RevisionApplicationError('the original manuscript must never be overwritten')
    if sha256_file(source) != expected_source_hash:
        raise RevisionApplicationError('source document SHA-256 mismatch')

    validated = [
        RevisionDraft.model_validate(
            item.model_dump(mode='python') if isinstance(item, RevisionDraft) else item
        ) for item in drafts
    ]
    if any(not _approved(draft) for draft in validated):
        raise RevisionApplicationError('only explicitly exact-text-approved drafts may be applied')
    if any(draft.manual_handling_required for draft in validated):
        raise RevisionApplicationError('manual-handling drafts cannot be applied automatically')
    _validate_conflicts(validated)

    structure = read_manuscript_structure(source)
    aliases, order = _structure_aliases(structure)
    table_ids = {
        element.table_index: element.element_id
        for element in structure.elements
        if element.element_type == 'table'
    }
    document = Document(source)
    paragraphs = _paragraph_map(document)

    targets: dict[str, tuple[Any, Paragraph]] = {}
    for draft in validated:
        if len(draft.target_element_ids) != 1:
            raise RevisionApplicationError('Phase 5 supports exactly one target per draft')
        target_id = draft.target_element_ids[0]
        element = aliases.get(target_id)
        if element is None or not element.paragraph_id:
            raise RevisionApplicationError(f'unknown or unsupported target element: {target_id}')
        paragraph = paragraphs.get(element.paragraph_id)
        if paragraph is None:
            raise RevisionApplicationError(f'target paragraph cannot be located: {target_id}')
        if draft.original_text_snapshot != element.text:
            raise RevisionApplicationError(f'stale target text snapshot: {draft.draft_id}')
        if draft.original_text_hash != sha256_text(element.text):
            raise RevisionApplicationError(f'stale target hash: {draft.draft_id}')
        reasons = complex_content_reasons(paragraph)
        if reasons:
            raise RevisionApplicationError(
                f'{draft.draft_id} requires manual handling: ' + ', '.join(reasons)
            )
        if draft.operation is RevisionOperation.REPLACE_HEADING and element.element_type != 'heading':
            raise RevisionApplicationError(f'{draft.draft_id} target is not an exact heading')
        if (
            draft.operation is RevisionOperation.REPLACE_FIGURE_CAPTION
            and element.element_type != 'figure_caption'
        ):
            raise RevisionApplicationError(f'{draft.draft_id} target is not a figure caption')
        if (
            draft.operation is RevisionOperation.REPLACE_TABLE_CAPTION
            and element.element_type != 'table_caption'
        ):
            raise RevisionApplicationError(f'{draft.draft_id} target is not a table caption')
        if draft.operation is RevisionOperation.REPLACE_TABLE_CELL:
            if element.element_type != 'table_cell_paragraph':
                raise RevisionApplicationError(f'{draft.draft_id} target is not a table cell')
            if (
                element.row_index != draft.table_row
                or element.cell_index != draft.table_column
            ):
                raise RevisionApplicationError(f'{draft.draft_id} table coordinates changed')
            if draft.table_id != table_ids.get(element.table_index):
                raise RevisionApplicationError(
                    f'{draft.draft_id} table ID changed')
        targets[draft.draft_id] = (element, paragraph)

    ordered = sorted(
        validated,
        key=lambda item: (
            order[item.target_element_ids[0]],
            item.operation.value,
            item.draft_id,
        ),
    )
    mutations: list[AppliedMutation] = []
    for index, draft in enumerate(ordered, start=1):
        element, paragraph = targets[draft.draft_id]
        old_text = element.text
        new_text = draft.text_for_application
        change_id = f'CHG-{index:04d}'
        if draft.operation in {
            RevisionOperation.REPLACE_PARAGRAPH,
            RevisionOperation.REPLACE_HEADING,
            RevisionOperation.REPLACE_TABLE_CELL,
            RevisionOperation.REPLACE_FIGURE_CAPTION,
            RevisionOperation.REPLACE_TABLE_CAPTION,
        }:
            _replace_paragraph_text(paragraph, new_text, draft, change_id)
        elif draft.operation is RevisionOperation.INSERT_BEFORE:
            _new_sibling(
                paragraph, new_text, before=True, draft=draft, change_id=change_id
            )
        elif draft.operation in {
            RevisionOperation.INSERT_AFTER,
            RevisionOperation.ADD_PARAGRAPH_TO_SECTION,
        }:
            _new_sibling(
                paragraph, new_text, before=False, draft=draft, change_id=change_id
            )
        elif draft.operation is RevisionOperation.DELETE_PARAGRAPH:
            _delete_paragraph(paragraph)
            new_text = ''
        else:
            raise RevisionApplicationError(f'unsupported operation: {draft.operation.value}')
        mutations.append(AppliedMutation(
            draft_id=draft.draft_id,
            action_id=draft.action_id,
            comment_ids=tuple(draft.comment_ids),
            operation=draft.operation.value,
            target_section=draft.target_section,
            target_element_id=draft.target_element_ids[0],
            old_text=old_text,
            new_text=new_text,
            highlight=draft.highlight.value,
        ))

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return tuple(mutations)
