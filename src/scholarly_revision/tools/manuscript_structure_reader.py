'''Read manuscript structure from DOCX without modifying or paginating it.'''

from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml.etree import XMLSyntaxError


class ManuscriptStructureError(ValueError):
    '''Raised when manuscript structure cannot be read deterministically.'''


@dataclass(frozen=True, slots=True)
class ManuscriptElement:
    order_index: int
    element_id: str
    element_type: str
    text: str
    paragraph_id: str | None = None
    section_id: str | None = None
    heading_level: int | None = None
    caption: str | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    citation_patterns: tuple[str, ...] = ()
    highlight_colors: tuple[str, ...] = ()
    page_number: int | None = None
    uncertain: bool = False
    uncertainty_reason: str | None = None

    def to_dict(self) -> dict[str, object]:
        return asdict(self)


@dataclass(frozen=True, slots=True)
class ManuscriptStructure:
    source_file: str
    title: str | None
    title_source: str | None
    title_uncertain: bool
    section_order: tuple[str, ...]
    outline: tuple[dict[str, object], ...]
    elements: tuple[ManuscriptElement, ...]
    reference_section_boundary: dict[str, object] | None
    explicit_page_count: int | None
    warnings: tuple[str, ...] = field(default_factory=tuple)

    def to_dict(self) -> dict[str, object]:
        payload = asdict(self)
        payload['elements'] = [element.to_dict() for element in self.elements]
        return payload


_HEADING_STYLE = re.compile(r'^heading\s+(?P<level>[1-9])$', re.IGNORECASE)
_FIGURE_CAPTION = re.compile(
    r'^\s*(?:figure|fig\.)\s*(?P<number>\d+|[ivxlcdm]+)?\s*[:.\-]?',
    re.IGNORECASE,
)
_TABLE_CAPTION = re.compile(
    r'^\s*(?:table|tbl\.)\s*(?P<number>\d+|[ivxlcdm]+)?\s*[:.\-]?',
    re.IGNORECASE,
)
_EQUATION_CAPTION = re.compile(
    r'^\s*(?:equation|eq\.)\s*(?P<number>\d+)?\s*[:.\-]?',
    re.IGNORECASE,
)
_EQUATION_NUMBER = re.compile(r'\(\s*\d+\s*\)\s*$')
_CITATION = re.compile(r'\[(?:\s*\d+\s*(?:[-,;]\s*\d+\s*)*)\]')
_REFERENCE_HEADING = re.compile(
    r'^\s*(?:references|bibliography|reference\s+list)\s*$',
    re.IGNORECASE,
)


def _style_name(paragraph: Paragraph) -> str | None:
    try:
        return paragraph.style.name if paragraph.style is not None else None
    except (AttributeError, KeyError):
        return None


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = _style_name(paragraph)
    match = _HEADING_STYLE.fullmatch(style_name or '')
    if match:
        return int(match.group('level'))
    properties = paragraph._p.pPr
    if properties is not None:
        outline = properties.find(qn('w:outlineLvl'))
        if outline is not None:
            raw = outline.get(qn('w:val'))
            if raw is not None and raw.isdigit():
                return int(raw) + 1
    return None


def _highlight_colors(paragraph: Paragraph) -> tuple[str, ...]:
    colors: list[str] = []
    for run in paragraph.runs:
        value = run.font.highlight_color
        if value is None:
            continue
        if isinstance(value, WD_COLOR_INDEX):
            name = value.name
        else:
            name = str(value)
        if name not in colors:
            colors.append(name)
    return tuple(colors)


def _contains_math(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]'))


def _looks_like_uncertain_heading(paragraph: Paragraph, text: str) -> bool:
    if not text or len(text) > 100 or text.endswith(('.', '?', '!')):
        return False
    visible_runs = [run for run in paragraph.runs if run.text.strip()]
    if visible_runs and all(bool(run.bold) for run in visible_runs):
        return True
    style = (_style_name(paragraph) or '').lower()
    return 'heading' in style or 'section' in style


def _explicit_page_count(source: Path) -> int | None:
    try:
        with ZipFile(source) as archive:
            raw = archive.read('docProps/app.xml')
    except (KeyError, OSError, BadZipFile):
        return None
    try:
        root = ElementTree.fromstring(raw)
    except ElementTree.ParseError:
        return None
    for element in root.iter():
        if element.tag.rsplit('}', 1)[-1] == 'Pages':
            value = (element.text or '').strip()
            if value.isdigit() and int(value) > 0:
                return int(value)
    return None


def _load_document(source: Path):
    if not source.exists():
        raise FileNotFoundError(f'manuscript DOCX not found: {source}')
    if not source.is_file():
        raise ManuscriptStructureError(f'manuscript path is not a file: {source}')
    if source.suffix.lower() != '.docx':
        raise ManuscriptStructureError('manuscript structure intake requires DOCX')
    try:
        return Document(source)
    except PermissionError as exc:
        raise ManuscriptStructureError(f'manuscript DOCX is unreadable: {source}') from exc
    except (PackageNotFoundError, BadZipFile, XMLSyntaxError, KeyError, ValueError) as exc:
        raise ManuscriptStructureError(f'invalid manuscript DOCX: {source}') from exc
    except OSError as exc:
        raise ManuscriptStructureError(f'unable to read manuscript DOCX: {source}') from exc


def read_manuscript_structure(path: str | Path) -> ManuscriptStructure:
    '''Extract structural facts only; never infer page or line locations.'''

    source = Path(path).expanduser().resolve()
    document = _load_document(source)
    elements: list[ManuscriptElement] = []
    outline: list[dict[str, object]] = []
    section_order: list[str] = []
    warnings: list[str] = []
    counters = {'PAR': 0, 'SEC': 0, 'TBL': 0, 'FIG': 0, 'EQ': 0, 'REF': 0}
    order_index = 0
    current_section: str | None = None
    in_references = False
    reference_boundary: dict[str, object] | None = None
    pending_table_caption: str | None = None
    first_nonempty: tuple[str, str | None] | None = None

    def next_id(prefix: str, width: int = 3) -> str:
        counters[prefix] += 1
        return f'{prefix}-{counters[prefix]:0{width}d}'

    def add_paragraph(
        paragraph: Paragraph,
        *,
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
    ) -> None:
        nonlocal order_index, current_section, in_references
        nonlocal reference_boundary, pending_table_caption, first_nonempty
        text = paragraph.text
        stripped = text.strip()
        paragraph_id = next_id('PAR', 4)
        style_name = _style_name(paragraph)
        level = _heading_level(paragraph)
        structural_id = paragraph_id
        element_type = 'table_cell_paragraph' if table_index is not None else 'paragraph'
        uncertain = False
        reason: str | None = None

        if stripped and first_nonempty is None:
            first_nonempty = (stripped, style_name)

        if table_index is None and _REFERENCE_HEADING.fullmatch(stripped):
            structural_id = next_id('REF')
            element_type = 'reference_section_heading'
            level = level or 1
            current_section = structural_id
            in_references = True
            reference_boundary = {
                'element_id': structural_id,
                'paragraph_id': paragraph_id,
                'order_index': order_index,
                'heading': stripped,
            }
            section_order.append(stripped)
            outline.append(
                {'section_id': structural_id, 'paragraph_id': paragraph_id,
                 'title': stripped, 'level': level, 'order_index': order_index}
            )
        elif table_index is None and level is not None:
            structural_id = next_id('SEC')
            element_type = 'heading'
            current_section = structural_id
            section_order.append(stripped)
            outline.append(
                {'section_id': structural_id, 'paragraph_id': paragraph_id,
                 'title': stripped, 'level': level, 'order_index': order_index}
            )
        elif table_index is None and _FIGURE_CAPTION.match(stripped):
            structural_id = next_id('FIG')
            element_type = 'figure_caption'
            if _FIGURE_CAPTION.match(stripped).group('number') is None:
                uncertain = True
                reason = 'Figure caption has no explicit figure number.'
        elif table_index is None and (
            _EQUATION_CAPTION.match(stripped)
            or (style_name or '').lower() == 'equation'
            or _contains_math(paragraph)
            or (_EQUATION_NUMBER.search(stripped) and any(char in stripped for char in '=+-/'))
        ):
            structural_id = next_id('EQ')
            element_type = 'equation'
            if not _contains_math(paragraph):
                uncertain = True
                reason = 'Equation identified from caption, numbering, or placeholder text.'
        elif in_references and stripped:
            structural_id = next_id('REF')
            element_type = 'reference_entry'
        elif table_index is None and _looks_like_uncertain_heading(paragraph, stripped):
            element_type = 'possible_heading'
            uncertain = True
            reason = 'Heading-like formatting lacks an explicit heading level.'

        if table_index is None and _TABLE_CAPTION.match(stripped):
            element_type = 'table_caption'
            pending_table_caption = stripped
        elif table_index is None and stripped and element_type not in {'table_caption'}:
            pending_table_caption = None

        elements.append(
            ManuscriptElement(
                order_index=order_index,
                element_id=structural_id,
                element_type=element_type,
                text=text,
                paragraph_id=paragraph_id,
                section_id=current_section,
                heading_level=level,
                caption=text if element_type.endswith('caption') else None,
                table_index=table_index,
                row_index=row_index,
                cell_index=cell_index,
                citation_patterns=tuple(_CITATION.findall(text)),
                highlight_colors=_highlight_colors(paragraph),
                uncertain=uncertain,
                uncertainty_reason=reason,
            )
        )
        order_index += 1

    table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn('w:p'):
            add_paragraph(Paragraph(child, document))
        elif child.tag == qn('w:tbl'):
            table_id = next_id('TBL')
            uncertain = pending_table_caption is None
            elements.append(
                ManuscriptElement(
                    order_index=order_index,
                    element_id=table_id,
                    element_type='table',
                    text='',
                    section_id=current_section,
                    caption=pending_table_caption,
                    table_index=table_index,
                    uncertain=uncertain,
                    uncertainty_reason=(
                        'No immediately preceding table caption was detected.'
                        if uncertain else None
                    ),
                )
            )
            order_index += 1
            table = Table(child, document)
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        add_paragraph(
                            paragraph,
                            table_index=table_index,
                            row_index=row_index,
                            cell_index=cell_index,
                        )
            table_index += 1
            pending_table_caption = None

    if not any(element.text.strip() or element.element_type == 'table' for element in elements):
        raise ManuscriptStructureError('manuscript DOCX contains no readable structure')

    title: str | None = None
    title_source: str | None = None
    title_uncertain = False
    core_title = (document.core_properties.title or '').strip()
    title_paragraph = next(
        (element for element in elements
         if element.paragraph_id and element.text.strip()
         and (_style_name(document.paragraphs[int(element.paragraph_id.split('-')[1]) - 1])
              if int(element.paragraph_id.split('-')[1]) <= len(document.paragraphs) else None) == 'Title'),
        None,
    )
    if title_paragraph is not None:
        title = title_paragraph.text.strip()
        title_source = 'paragraph_style'
    elif core_title:
        title = core_title
        title_source = 'core_properties'
    elif first_nonempty is not None and not _HEADING_STYLE.fullmatch(first_nonempty[1] or ''):
        title = first_nonempty[0]
        title_source = 'first_nonempty_paragraph'
        title_uncertain = True
        warnings.append('Document title was inferred only as a candidate from the first paragraph.')

    if any(element.uncertain for element in elements):
        warnings.append('One or more structural elements require manual review.')
    return ManuscriptStructure(
        source_file=source.name,
        title=title,
        title_source=title_source,
        title_uncertain=title_uncertain,
        section_order=tuple(section_order),
        outline=tuple(outline),
        elements=tuple(elements),
        reference_section_boundary=reference_boundary,
        explicit_page_count=_explicit_page_count(source),
        warnings=tuple(warnings),
    )
