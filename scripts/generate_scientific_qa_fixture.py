'''Generate the anonymous synthetic Phase 6 DOCX fixture.'''
from __future__ import annotations
import sys
from pathlib import Path
ROOT=Path(__file__).resolve().parents[1];SOURCE=ROOT/'src'
if str(SOURCE) not in sys.path:sys.path.insert(0,str(SOURCE))
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from scholarly_revision.tools.docx_highlight_manager import apply_revision_highlight

def generate(path:Path)->Path:
    document=Document();document.core_properties.title='Anonymous Synthetic Audit Study'
    section=document.sections[0];section.header.paragraphs[0].text='Old Journal Running Header 2019'
    document.add_heading('Anonymous Synthetic Audit Study',0)
    document.add_paragraph('FirstName LastName; Affiliation placeholder; email@example.com')
    document.add_heading('Abstract',1)
    document.add_paragraph('The final synthetic_score = 8.0 units is reported as RES-0001 [1].')
    document.add_heading('Introduction',1)
    document.add_paragraph('The adaptive-mode procedure uses an anonymous placeholder setting [2, 3].')
    document.add_heading('Results',1)
    document.add_paragraph('The synthetic_score = 9.0 units in the same setting [1, 1].')
    document.add_paragraph('The value increased from 40 to 50, a 20% improvement.')
    document.add_paragraph('We conducted an experiment with no evidence record.')
    document.add_paragraph('The difference was statistically significant (p = 0.04).')
    document.add_paragraph('A malformed citation remains [4-3].')
    document.add_paragraph('As shown in Fig. 2, the anonymous process is illustrated.')
    document.add_paragraph('Figure 1. Anonymous synthetic process.',style='Caption')
    document.add_paragraph('Table 1. Anonymous synthetic values.',style='Caption')
    table=document.add_table(rows=2,cols=2);table.cell(0,0).text='Metric';table.cell(0,1).text='Value'
    table.cell(1,0).text='synthetic_score';table.cell(1,1).text='9.0 units'
    document.add_paragraph('Table 1. Anonymous synthetic values.',style='Caption')
    document.add_heading('Notation',1)
    document.add_paragraph('y = a*x (1)')
    document.add_paragraph('z = b*x (1)')
    document.add_paragraph('where x is the anonymous input.')
    document.add_paragraph('Later, x denotes the anonymous output.')
    document.add_paragraph('The adaptive method form is also used.')
    document.add_heading('References',1)
    p=document.add_paragraph();run=p.add_run('[1] A. Example. Anonymous reference title one. Synthetic Source. 2020. 10.1234/syn.001')
    apply_revision_highlight(run,'BRIGHT_GREEN',change_id='CHG-0001')
    document.add_paragraph('[3] B. Example. Anonymous reference title three. Synthetic Source. 2021. 10.1234/syn.003')
    document.add_paragraph('[3] C. Example. Anonymous duplicate number title. Synthetic Source. 2022. 10.1234/syn.003')
    document.add_paragraph('[4] D. Example. Anonymous uncited title. Synthetic Source. 2023.')
    path.parent.mkdir(parents=True,exist_ok=True);document.save(path);return path

if __name__=='__main__':
    target=ROOT/'tests'/'fixtures'/'synthetic_scientific_audit_manuscript.docx'
    generate(target);print(target)
