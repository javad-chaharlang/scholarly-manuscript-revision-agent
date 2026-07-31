from pathlib import Path

from docx import Document

from scholarly_revision.models.response_package import ResponsePackage
from scholarly_revision.tools.response_docx_builder import build_response_docx


def test_response_docx_generation_is_editable_and_preserves_comment(tmp_path: Path) -> None:
    payload = {
        'generated_at': '2030-01-01T00:00:00Z',
        'manuscript_title': 'Anonymous', 'manuscript_id': 'SYN',
        'journal': 'Synthetic', 'revision_round': 1,
        'cover_letter': {'body_paragraphs': ['Verified metadata only.']},
        'sections': [{
            'section_id': 'R1', 'title': 'Reviewer 1',
            'reviewer_source': 'REVIEWER', 'reviewer_number': 1,
            'entries': [{
                'response_entry_id': 'RESP-0001', 'reviewer_source': 'REVIEWER',
                'reviewer_number': 1, 'comment_id': 'R1-C01',
                'sequence_number': 1, 'exact_comment': 'Exact synthetic comment.',
                'author_response': 'Direct response.', 'highlight': 'YELLOW',
                'response_status': 'APPROVED', 'location_status': 'NOT_REQUIRED',
                'evidence_status': 'NOT_REQUIRED', 'author_approved': True,
                'resolution': 'NOT_APPLICABLE'
            }]
        }], 'closing_statement': 'Closing.'
    }
    path = build_response_docx(ResponsePackage.model_validate(payload), tmp_path / 'response.docx')
    document = Document(path)
    assert document.paragraphs[0].text == 'Response to the Editor and Reviewers'
    assert sum('Exact synthetic comment.' in cell.text
               for table in document.tables for row in table.rows for cell in row.cells) == 1
