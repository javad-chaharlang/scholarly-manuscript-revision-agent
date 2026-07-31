from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from scholarly_revision.tools.docx_clean_copy import (
    create_clean_copy,
    validate_text_equivalence,
)
from scholarly_revision.tools.docx_highlight_manager import apply_revision_highlight


def test_removes_only_system_highlights_and_preserves_author_highlight(tmp_path: Path) -> None:
    highlighted = tmp_path / 'highlighted.docx'
    clean = tmp_path / 'clean.docx'
    document = Document()
    paragraph = document.add_paragraph()
    author = paragraph.add_run('Author highlight. ')
    author.font.highlight_color = WD_COLOR_INDEX.YELLOW
    system = paragraph.add_run('System highlight.')
    apply_revision_highlight(system, 'VIOLET', change_id='CHG-0001')
    document.save(highlighted)
    _, removed = create_clean_copy(highlighted, clean)
    assert removed == 1
    assert validate_text_equivalence(highlighted, clean)
    runs = Document(clean).paragraphs[0].runs
    assert runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW
    assert runs[1].font.highlight_color is None
