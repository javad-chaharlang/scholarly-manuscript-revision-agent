from pathlib import Path

from docx import Document

from scholarly_revision.tools.location_verifier import verify_location


def manuscript(tmp_path: Path) -> Path:
    path = tmp_path / 'locations.docx'
    document = Document()
    document.add_heading('2.3 Methods', 1)
    document.add_paragraph('Stable paragraph.')
    document.add_paragraph('Table 8. Synthetic object.', style='Caption')
    document.add_table(rows=1, cols=1).cell(0, 0).text = 'Cell'
    document.add_paragraph('Equation placeholder (14)')
    document.add_heading('References', 1)
    document.add_paragraph('[29] Anonymous synthetic record.')
    document.save(path)
    return path


def test_section_and_object_locations_verify(tmp_path: Path) -> None:
    path = manuscript(tmp_path)
    assert verify_location(path, 'Section 2.3').verified
    assert verify_location(path, 'Table 8').verified
    assert verify_location(path, 'Equation (14)').verified
    assert verify_location(path, 'Reference [29]').verified


def test_page_and_line_location_requires_explicit_render_metadata(tmp_path: Path) -> None:
    path = manuscript(tmp_path)
    result = verify_location(path, 'Page 2, Lines 3-5')
    assert not result.verified
    metadata = {'verified_locations': ['Page 2, Lines 3-5']}
    assert verify_location(
        path, 'Page 2, Lines 3-5', page_metadata=metadata
    ).status.value == 'PAGE_AND_LINES_VERIFIED'
