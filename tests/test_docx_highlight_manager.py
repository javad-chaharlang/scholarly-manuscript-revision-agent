from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from scholarly_revision.tools.docx_highlight_manager import (
    apply_revision_highlight,
    audit_revision_highlights,
)


def test_exact_highlight_mapping_and_audit(tmp_path: Path) -> None:
    path = tmp_path / 'highlights.docx'
    document = Document()
    paragraph = document.add_paragraph()
    apply_revision_highlight(
        paragraph.add_run('Reviewer one'), 'YELLOW', change_id='CHG-0001'
    )
    apply_revision_highlight(
        paragraph.add_run('Reviewer two'), 'BRIGHT_GREEN', change_id='CHG-0002'
    )
    apply_revision_highlight(
        paragraph.add_run('Shared'), 'VIOLET', change_id='CHG-0003'
    )
    document.save(path)
    audit = audit_revision_highlights(path)
    assert audit['passed']
    assert [item['highlight'] for item in audit['records']] == [
        'YELLOW', 'BRIGHT_GREEN', 'VIOLET'
    ]
