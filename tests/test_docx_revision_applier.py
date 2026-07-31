from datetime import UTC, datetime
from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement

from phase5_helpers import MANUSCRIPT
from scholarly_revision.models.revision_draft import RevisionDraft
from scholarly_revision.models.enums import (
    RevisionDraftStatus,
    RevisionOperation,
    RevisionTextApprovalState,
)
from scholarly_revision.services.project_workspace import sha256_file
from scholarly_revision.services.revision_drafting_service import sha256_text
from scholarly_revision.tools.docx_revision_applier import (
    RevisionApplicationError,
    apply_docx_revisions,
)
from scholarly_revision.tools.manuscript_structure_reader import read_manuscript_structure


def approved_draft(target='PAR-0006', text='Synthetic replacement.', draft_id='DRAFT-0001'):
    structure = read_manuscript_structure(MANUSCRIPT)
    element = next(
        item for item in structure.elements
        if target in {item.element_id, item.paragraph_id}
    )
    now = datetime(2026, 1, 1, tzinfo=UTC)
    return RevisionDraft(
        draft_id=draft_id, action_id=f'ACT-{draft_id[-4:]}',
        comment_ids=['R1-C01'], source_document_hash=sha256_file(MANUSCRIPT),
        target_element_ids=[target], target_section='Synthetic',
        operation='REPLACE_PARAGRAPH', original_text_snapshot=element.text,
        original_text_hash=sha256_text(element.text), proposed_text=text,
        drafting_rationale='Synthetic.', highlight='YELLOW',
        draft_status='APPROVED', approval_state='APPROVED',
        author_decision='APPROVE_TEXT', approved_text=text,
        decision_maker='author', decision_timestamp=now,
        created_at=now, updated_at=now,
    )


def test_exact_application_preserves_style_and_source(tmp_path: Path) -> None:
    before = sha256(MANUSCRIPT.read_bytes()).hexdigest()
    output = tmp_path / 'revised.docx'
    draft = approved_draft(target='SEC-004', text='Revised Synthetic Heading')
    draft = draft.model_copy(update={'operation': RevisionOperation.REPLACE_HEADING})
    apply_docx_revisions(MANUSCRIPT, output, [draft], expected_source_hash=before)
    assert sha256(MANUSCRIPT.read_bytes()).hexdigest() == before
    paragraph = next(p for p in Document(output).paragraphs if p.text == 'Revised Synthetic Heading')
    assert paragraph.style.name == 'Heading 1'


def test_stale_overlap_and_unapproved_fail_safely(tmp_path: Path) -> None:
    stale = approved_draft().model_copy(update={'original_text_hash': 'f' * 64})
    with pytest.raises(RevisionApplicationError, match='stale target hash'):
        apply_docx_revisions(MANUSCRIPT, tmp_path / 'stale.docx', [stale],
                             expected_source_hash=sha256_file(MANUSCRIPT))
    first = approved_draft(draft_id='DRAFT-0001')
    second = approved_draft(draft_id='DRAFT-0002')
    with pytest.raises(RevisionApplicationError, match='overlapping'):
        apply_docx_revisions(MANUSCRIPT, tmp_path / 'overlap.docx', [first, second],
                             expected_source_hash=sha256_file(MANUSCRIPT))
    unapproved = first.model_copy(update={
        'draft_status': RevisionDraftStatus.AWAITING_TEXT_APPROVAL,
        'approval_state': RevisionTextApprovalState.PENDING,
        'author_decision': None, 'approved_text': None,
    })
    with pytest.raises(RevisionApplicationError, match='only explicitly'):
        apply_docx_revisions(MANUSCRIPT, tmp_path / 'unapproved.docx', [unapproved],
                             expected_source_hash=sha256_file(MANUSCRIPT))


def test_complex_field_target_is_blocked(tmp_path: Path) -> None:
    source = tmp_path / 'field.docx'
    document = Document()
    paragraph = document.add_paragraph('Synthetic field target.')
    field = OxmlElement('w:fldSimple')
    paragraph._p.append(field)
    document.save(source)
    now = datetime(2026, 1, 1, tzinfo=UTC)
    draft = RevisionDraft(
        draft_id='DRAFT-0001', action_id='ACT-0001', comment_ids=['R1-C01'],
        source_document_hash=sha256_file(source), target_element_ids=['PAR-0001'],
        target_section='Synthetic', operation='REPLACE_PARAGRAPH',
        original_text_snapshot='Synthetic field target.',
        original_text_hash=sha256_text('Synthetic field target.'),
        proposed_text='Replacement.', drafting_rationale='Synthetic.',
        highlight='YELLOW', draft_status='APPROVED', approval_state='APPROVED',
        author_decision='APPROVE_TEXT', approved_text='Replacement.',
        created_at=now, updated_at=now,
    )
    with pytest.raises(RevisionApplicationError, match='manual handling'):
        apply_docx_revisions(source, tmp_path / 'blocked.docx', [draft],
                             expected_source_hash=sha256_file(source))
