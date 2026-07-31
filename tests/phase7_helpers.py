from __future__ import annotations

from pathlib import Path

from docx import Document

from scholarly_revision.models.enums import ResultStatus
from scholarly_revision.models.project import OutputNames, ProjectManifest
from scholarly_revision.models.release import MANUAL_VISUAL_QA_ARTIFACTS
from scholarly_revision.models.reviewer import ReviewerComment
from scholarly_revision.services.config_loader import save_project_manifest
from scholarly_revision.services.gap_analysis_service import write_json
from scholarly_revision.services.manual_visual_qa_service import (
    import_manual_visual_qa_decisions,
)
from scholarly_revision.services.project_workspace import sha256_file
from scholarly_revision.tools.workbook_builder import build_revision_workbook
from scholarly_revision.workflows.finalization_workflow import (
    generate_response_letter, run_final_consistency_check, verify_response_letter,
)
from scholarly_revision.workflows.scientific_qa_workflow import (
    run_scientific_qa_workflow,
)


def approved_manual_visual_qa_payload(root: Path) -> dict:
    return {
        'schema_version': 1,
        'decisions': [{
            'artifact_name': name,
            'artifact_sha256': sha256_file(root / 'outputs' / name),
            'opened_successfully': True,
            'repair_warning_present': False,
            'layout_acceptable': True,
            'highlights_verified': True,
            'tables_and_captions_acceptable': True,
            'clean_highlight_text_equivalence_confirmed': True,
            'reviewer_notes': 'Explicit anonymous synthetic inspection decision.',
            'decision_maker': 'anonymous-visual-reviewer',
            'decision_timestamp': '2030-01-01T00:00:00Z',
            'decision': 'APPROVED',
        } for name in MANUAL_VISUAL_QA_ARTIFACTS],
    }


def make_ready_phase7_project(tmp_path: Path) -> Path:
    root = tmp_path / 'anonymous-phase7-project'
    for name in ('working', 'outputs', 'audit', 'config', 'rendered', 'input'):
        (root / name).mkdir(parents=True, exist_ok=True)
    manifest = ProjectManifest(
        project_name='Anonymous Phase Seven',
        manuscript_id='SYNTHETIC-PHASE-7',
        manuscript_title='Anonymous Finalization Fixture',
        journal='Synthetic Journal',
        revision_round=1,
        manuscript_language='English',
        response_language='English',
        citation_style='none',
        reviewer_count=1,
        result_status=ResultStatus.NOT_APPLICABLE,
        output_names=OutputNames(
            highlighted_manuscript='Revised_Manuscript_Highlighted.docx',
            clean_manuscript='Revised_Manuscript_Clean.docx',
            revision_workbook='Revision_Master.xlsx',
            response_letter='Response_to_Reviewers.docx',
            qa_report='Final_QA_Report.xlsx',
            audit_log='audit.json',
        ),
    )
    save_project_manifest(manifest, root / 'config' / 'project_manifest.yaml')
    comment = ReviewerComment(
        comment_id='R1-C01', reviewer_source='REVIEWER', reviewer_number=1,
        sequence_number=1,
        original_comment='Please clarify the anonymous fixture scope.',
    )
    write_json(
        root / 'working' / 'reviewer_comments.json',
        [comment.model_dump(mode='json')],
    )
    write_json(root / 'working' / 'revision_plan.json', {
        'schema_version': 1, 'approval_gate_status': 'APPROVED', 'actions': [],
    })
    write_json(root / 'audit' / 'change_log.json', {
        'schema_version': 1, 'changes': [],
    })
    manuscript = Document()
    manuscript.add_heading('Anonymous Finalization Fixture', 0)
    manuscript.add_heading('Introduction', 1)
    manuscript.add_paragraph(
        'This anonymous fixture contains only plain text and no scientific claims.'
    )
    highlighted = root / 'outputs' / 'Revised_Manuscript_Highlighted.docx'
    clean = root / 'outputs' / 'Revised_Manuscript_Clean.docx'
    manuscript.save(highlighted)
    manuscript.save(clean)
    build_revision_workbook(
        root / 'outputs' / 'Revision_Master.xlsx', [comment]
    )
    draft = root / 'working' / 'completed_response.json'
    write_json(draft, {'entries': [{
        'response_entry_id': 'RESP-0001',
        'reviewer_source': 'REVIEWER',
        'reviewer_number': 1,
        'comment_id': 'R1-C01',
        'sequence_number': 1,
        'exact_comment': comment.original_comment,
        'author_response': (
            'The author confirms that the file is an anonymous workflow fixture '
            'and does not present a scientific study.'
        ),
        'changes_made': '',
        'verified_locations': [],
        'related_action_ids': [],
        'related_change_ids': [],
        'related_evidence_ids': [],
        'related_reference_ids': [],
        'highlight': 'YELLOW',
        'response_status': 'APPROVED',
        'location_status': 'NOT_REQUIRED',
        'evidence_status': 'NOT_REQUIRED',
        'author_approved': True,
        'verification_notes': [],
        'resolution': 'RESPECTFULLY_DECLINED',
        'approved_interpretation': 'The request concerns fixture scope.',
        'unresolved_limitations': [],
        'author_justification': (
            'No manuscript revision is applicable because the document is a '
            'non-scientific anonymous workflow fixture.'
        ),
    }]})
    generated = generate_response_letter(root, draft)
    verification = verify_response_letter(root, generated.response_letter_path)
    assert verification.passed
    run_scientific_qa_workflow(
        project_root=root,
        highlighted_manuscript=highlighted,
        clean_manuscript=clean,
    )
    import_manual_visual_qa_decisions(
        root, approved_manual_visual_qa_payload(root)
    )
    write_json(root / 'audit' / 'final_release_approval.json', {
        'approved': True,
        'decision_maker': 'anonymous-author',
        'decision_timestamp': '2030-01-01T00:00:00Z',
        'scope': 'anonymous synthetic release fixture',
    })
    result = run_final_consistency_check(root)
    assert result.readiness == 'READY'
    return root
