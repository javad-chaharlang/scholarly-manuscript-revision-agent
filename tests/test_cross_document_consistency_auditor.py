from pathlib import Path

from docx import Document

from phase7_helpers import make_ready_phase7_project
from scholarly_revision.models.response_package import (
    ResponseEntry, ResponsePackage, ResponseStatus,
)
from scholarly_revision.services.gap_analysis_service import read_json
from scholarly_revision.tools.cross_document_consistency_auditor import (
    audit_cross_document_consistency,
)


def package(root: Path) -> ResponsePackage:
    return ResponsePackage.model_validate(
        read_json(root / 'working' / 'response_package.json')
    )


def test_corrected_scenario_has_no_consistency_findings(tmp_path: Path) -> None:
    root = make_ready_phase7_project(tmp_path)
    result = audit_cross_document_consistency(root)
    assert result.passed
    assert not result.findings


def test_missing_response_and_clean_text_mismatch_block(tmp_path: Path) -> None:
    root = make_ready_phase7_project(tmp_path)
    source = package(root)
    missing = ResponsePackage.model_validate({
        **source.model_dump(mode='python'),
        'sections': [], 'package_status': ResponseStatus.BLOCKED,
    })
    document = Document(root / 'outputs' / 'Revised_Manuscript_Clean.docx')
    document.add_paragraph('Intentional anonymous mismatch.')
    document.save(root / 'outputs' / 'Revised_Manuscript_Clean.docx')
    result = audit_cross_document_consistency(root, missing)
    descriptions = ' '.join(item.description for item in result.findings)
    assert 'missing from the response' in descriptions
    assert 'does not match' in descriptions
    assert not result.passed


def test_false_numerical_revision_claim_is_detected(tmp_path: Path) -> None:
    root = make_ready_phase7_project(tmp_path)
    source = package(root)
    entry = ResponseEntry.model_validate({
        **source.entries[0].model_dump(mode='python'),
        'response_status': ResponseStatus.APPROVED,
        'changes_made': 'The value was revised to 99.9%.',
        'related_change_ids': ['CHG-9999'],
    })
    section = source.sections[0].model_copy(update={'entries': [entry]})
    changed = ResponsePackage.model_validate({
        **source.model_dump(mode='python'),
        'sections': [section], 'package_status': ResponseStatus.DRAFTED,
    })
    result = audit_cross_document_consistency(root, changed)
    categories = {item.category.value for item in result.findings}
    assert {'CHANGE_CLAIM', 'NUMERICAL'} <= categories


def test_response_docx_fields_must_match_package(tmp_path: Path) -> None:
    root = make_ready_phase7_project(tmp_path)
    path = root / 'outputs' / 'Response_to_Reviewers.docx'
    document = Document(path)
    response_table = next(table for table in document.tables if len(table.rows) == 5)
    response_table.rows[2].cells[0].text = (
        'Changes made in the manuscript:\nAltered change claim.'
    )
    document.save(path)
    result = audit_cross_document_consistency(root)
    assert 'CHANGE_CLAIM' in {
        item.category.value for item in result.findings
    }
    assert not result.passed
